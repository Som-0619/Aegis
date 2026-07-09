"""
Unit tests for the DocumentReader agent.
"""

import pytest
import os
from agents.document_reader import DocumentReader, DocumentReaderError

def test_document_reader_instantiation():
    """Verifies that DocumentReader can be instantiated successfully."""
    reader = DocumentReader()
    assert reader is not None

def test_document_reader_txt_read(tmp_path):
    """Verifies reading a standard text file."""
    reader = DocumentReader()
    
    # Create temp text file
    temp_file = tmp_path / "test_project.txt"
    temp_content = "Project Alpha status update: all good."
    temp_file.write_text(temp_content)
    
    parsed_text = reader.read_file(str(temp_file))
    assert temp_content in parsed_text

def test_document_reader_csv_read(tmp_path):
    """Verifies reading and parsing a CSV file."""
    reader = DocumentReader()
    
    # Create a temp CSV file
    temp_file = tmp_path / "test_data.csv"
    csv_content = "Task,Owner,Status\nDatabase Design,John,In Progress\nAPI Implementation,Jane,Not Started"
    temp_file.write_text(csv_content)
    
    parsed_text = reader.read_file(str(temp_file))
    
    assert "Task | Owner | Status" in parsed_text
    assert "Database Design | John | In Progress" in parsed_text
    assert "API Implementation | Jane | Not Started" in parsed_text

def test_document_reader_docx_read(tmp_path):
    """Verifies reading and parsing a DOCX file using python-docx library."""
    import docx
    reader = DocumentReader()
    
    # Create a temp DOCX file
    docx_path = tmp_path / "test_document.docx"
    doc = docx.Document()
    doc.add_paragraph("This is paragraph one of the project plan.")
    doc.add_paragraph("This is paragraph two.")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Phase"
    table.cell(0, 1).text = "Due Date"
    table.cell(1, 0).text = "Kickoff"
    table.cell(1, 1).text = "2026-07-08"
    
    doc.save(str(docx_path))
    
    parsed_text = reader.read_file(str(docx_path))
    
    assert "This is paragraph one of the project plan." in parsed_text
    assert "This is paragraph two." in parsed_text
    assert "Phase | Due Date" in parsed_text
    assert "Kickoff | 2026-07-08" in parsed_text

def test_document_reader_pdf_read(tmp_path):
    """Verifies reading and parsing a PDF file using PyMuPDF (fitz)."""
    import fitz
    reader = DocumentReader()
    
    # Create a temp PDF file
    pdf_path = tmp_path / "test_document.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Project Health Review: Project Omega")
    doc.save(str(pdf_path))
    doc.close()
    
    parsed_text = reader.read_file(str(pdf_path))
    
    assert "Project Health Review: Project Omega" in parsed_text
    assert "Page 1" in parsed_text

def test_document_reader_errors(tmp_path):
    """Verifies that DocumentReader raises descriptive DocumentReaderError exceptions on invalid inputs."""
    reader = DocumentReader()
    
    # 1. Non-existent file
    with pytest.raises(DocumentReaderError) as exc_info:
        reader.read_file("non_existent_file_path.pdf")
    assert "does not exist" in str(exc_info.value)
    
    # 2. Empty file path
    with pytest.raises(DocumentReaderError) as exc_info:
        reader.read_file("")
    assert "cannot be empty" in str(exc_info.value)
    
    # 3. Empty file (0 bytes)
    empty_file = tmp_path / "empty.txt"
    empty_file.write_text("")
    with pytest.raises(DocumentReaderError) as exc_info:
        reader.read_file(str(empty_file))
    assert "is empty" in str(exc_info.value)

def test_document_reader_excel_read(tmp_path):
    """Verifies reading and parsing an Excel sheet file."""
    import pandas as pd
    reader = DocumentReader()
    
    excel_path = tmp_path / "test_project.xlsx"
    df = pd.DataFrame({
        "Milestone": ["Design", "Beta"],
        "Status": ["Completed", "Delayed"]
    })
    # Save as Excel sheet using openpyxl
    df.to_excel(str(excel_path), sheet_name="Timeline", index=False)
    
    parsed_text = reader.read_file(str(excel_path))
    
    assert "Sheet: Timeline" in parsed_text
    assert "Milestone | Status" in parsed_text
    assert "Design | Completed" in parsed_text
    assert "Beta | Delayed" in parsed_text
