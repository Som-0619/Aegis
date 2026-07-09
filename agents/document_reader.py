"""
Document Reader Agent.
Handles reading and extracting raw text from various project document formats (PDF, DOCX, TXT, CSV).
No AI/LLM processing is applied here; it purely extracts raw text.
"""

import os
from typing import List
from utils.logger import get_logger

logger = get_logger(__name__)

class DocumentReaderError(Exception):
    """Custom exception raised when a document cannot be parsed or read."""
    pass

class DocumentReader:
    """Agent responsible for loading files and converting them into raw text."""
    
    def __init__(self):
        logger.info("DocumentReader agent initialized.")

    def read_file(self, filepath: str) -> str:
        """
        Reads a document based on its extension and returns the plain text contents.
        Auto-detects document types (.pdf, .docx, .txt, .csv, .md) and raises 
        DocumentReaderError instead of crashing on failures.
        
        Args:
            filepath: Path to the input file.
            
        Returns:
            Extracted clean plain text content.
            
        Raises:
            DocumentReaderError: If the file is missing, empty, or parsing fails.
        """
        if not filepath:
            raise DocumentReaderError("Filepath cannot be empty or None.")

        if not os.path.exists(filepath):
            logger.error(f"File not found: {filepath}")
            raise DocumentReaderError(f"The file '{filepath}' does not exist.")

        if os.path.getsize(filepath) == 0:
            logger.error(f"File is empty: {filepath}")
            raise DocumentReaderError(f"The file '{filepath}' is empty (0 bytes).")
            
        ext = os.path.splitext(filepath)[1].lower()
        logger.info(f"Detecting format: '{ext}' for file: '{filepath}'")
        
        try:
            if ext in (".txt", ".md"):
                return self._read_text(filepath)
            elif ext == ".pdf":
                return self._read_pdf(filepath)
            elif ext == ".docx":
                return self._read_docx(filepath)
            elif ext == ".csv":
                return self._read_csv(filepath)
            elif ext in (".xlsx", ".xls"):
                return self._read_excel(filepath)
            else:
                logger.warning(f"Unsupported extension '{ext}' for '{filepath}'. Attempting text fallback.")
                return self._read_text(filepath)
        except DocumentReaderError:
            # Re-raise custom exceptions directly
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing file '{filepath}': {e}")
            raise DocumentReaderError(f"Failed to parse document '{os.path.basename(filepath)}': {str(e)}")

    def _read_text(self, filepath: str) -> str:
        """Reads plain text / markdown files."""
        try:
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            logger.info(f"Successfully read TXT/MD file '{filepath}'. Char count: {len(content)}")
            return content
        except Exception as e:
            logger.error(f"Failed reading TXT file {filepath}: {e}")
            raise DocumentReaderError(f"Error reading plain text file: {str(e)}")

    def _read_pdf(self, filepath: str) -> str:
        """Reads PDF files using PyMuPDF (fitz) and extracts text page-by-page."""
        try:
            import fitz  # type: ignore
        except ImportError as e:
            logger.error("PyMuPDF (fitz) package is not installed.")
            raise DocumentReaderError("System error: PyMuPDF is not installed in the current environment.") from e

        try:
            text_pages: List[str] = []
            with fitz.open(filepath) as doc:
                page_count = doc.page_count
                if page_count == 0:
                    raise DocumentReaderError("PDF contains no pages.")
                for i, page in enumerate(doc):
                    page_text = page.get_text()
                    text_pages.append(f"--- Page {i + 1} ---\n{page_text}")
            
            full_text = "\n".join(text_pages).strip()
            if not full_text:
                logger.warning(f"PDF '{filepath}' has no extractable text content.")
                
            logger.info(f"Successfully read PDF file '{filepath}'. Page count: {page_count}")
            return full_text
        except Exception as e:
            logger.error(f"Error reading PDF {filepath}: {e}")
            raise DocumentReaderError(f"Error parsing PDF document: {str(e)}")

    def _read_docx(self, filepath: str) -> str:
        """Reads DOCX files using python-docx and extracts text from paragraphs and tables."""
        try:
            import docx  # type: ignore
        except ImportError as e:
            logger.error("python-docx package is not installed.")
            raise DocumentReaderError("System error: python-docx is not installed in the current environment.") from e

        try:
            doc = docx.Document(filepath)
            extracted_text = []

            # 1. Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text)

            # 2. Extract from tables
            for table_idx, table in enumerate(doc.tables):
                extracted_text.append(f"\n--- Table {table_idx + 1} ---")
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        # Clean up cells to remove immediate merges/duplicates
                        cleaned_row = []
                        for cell in row_cells:
                            if not cleaned_row or cleaned_row[-1] != cell:
                                cleaned_row.append(cell)
                        extracted_text.append(" | ".join(cleaned_row))
            
            full_text = "\n".join(extracted_text).strip()
            logger.info(f"Successfully read DOCX file '{filepath}'. Element block count: {len(extracted_text)}")
            return full_text
        except Exception as e:
            logger.error(f"Error reading DOCX {filepath}: {e}")
            raise DocumentReaderError(f"Error parsing DOCX document: {str(e)}")

    def _read_csv(self, filepath: str) -> str:
        """Reads CSV files using pandas and formats contents into a readable plain-text table."""
        try:
            import pandas as pd  # type: ignore
        except ImportError as e:
            logger.error("pandas package is not installed.")
            raise DocumentReaderError("System error: pandas is not installed in the current environment.") from e

        try:
            # Read CSV
            df = pd.read_csv(filepath)
            
            if df.empty:
                logger.warning(f"CSV file '{filepath}' is empty.")
                return ""

            # Format as standard markdown-like pipe representation
            lines = []
            
            # Header row
            headers = [str(col).strip() for col in df.columns]
            lines.append(" | ".join(headers))
            lines.append("-|-".join(["---"] * len(headers)))
            
            # Data rows
            for _, row in df.iterrows():
                row_vals = [str(val).strip() for val in row.values]
                lines.append(" | ".join(row_vals))
                
            full_text = "\n".join(lines).strip()
            logger.info(f"Successfully read CSV file '{filepath}'. Rows: {len(df)}")
            return full_text
        except Exception as e:
            logger.error(f"Error reading CSV {filepath}: {e}")
            raise DocumentReaderError(f"Error parsing CSV document: {str(e)}")

    def _read_excel(self, filepath: str) -> str:
        """Reads Excel (.xlsx, .xls) files using pandas and formats contents as markdown-like tables."""
        try:
            import pandas as pd  # type: ignore
        except ImportError as e:
            logger.error("pandas package is not installed.")
            raise DocumentReaderError("System error: pandas is not installed in the current environment.") from e

        try:
            # Read all sheets in the excel file
            xls = pd.ExcelFile(filepath)
            sheet_texts = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                # Drop fully empty rows/cols to keep text clean
                df = df.dropna(how="all").dropna(axis=1, how="all")
                if not df.empty:
                    # Format as standard markdown-like pipe representation
                    lines = []
                    headers = [str(col).strip() for col in df.columns]
                    lines.append(f"### Sheet: {sheet_name}")
                    lines.append(" | ".join(headers))
                    lines.append("-|-".join(["---"] * len(headers)))
                    
                    for _, row in df.iterrows():
                        row_vals = [str(val).strip() for val in row.values]
                        lines.append(" | ".join(row_vals))
                    sheet_texts.append("\n".join(lines))
            
            full_text = "\n\n".join(sheet_texts).strip()
            if not full_text:
                logger.warning(f"Excel file '{filepath}' contains no extractable data.")
            logger.info(f"Successfully read Excel file '{filepath}'. Sheet count: {len(xls.sheet_names)}")
            return full_text
        except Exception as e:
            logger.error(f"Error reading Excel {filepath}: {e}")
            raise DocumentReaderError(f"Error parsing Excel document: {str(e)}")
